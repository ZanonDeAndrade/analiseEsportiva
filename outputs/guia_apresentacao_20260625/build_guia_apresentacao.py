from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Microsoft User\Desktop\TrabalhoIa\AnaliseEsportiva")
OUT = ROOT / "outputs" / "guia_apresentacao_20260625"
DOCX = OUT / "guia_estudo_betintel_ai_apresentacao.docx"
PDF_PATH = Path(r"C:\Users\Microsoft User\Downloads")
EXCEL = ROOT / "outputs" / "worldcup_last5_20260625" / "betintel_world_cups_2006_2022_training_dataset.xlsx"


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "111827",
    "muted": "4B5563",
    "light_blue": "E8EEF5",
    "light_gray": "F3F4F6",
    "callout": "F4F6F9",
    "warn": "FFF7E6",
    "warn_border": "B7791F",
    "green_fill": "EAF7EA",
}


def load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


model = load_json(ROOT / "backend" / "artifacts" / "model.json")
evaluation = load_json(ROOT / "backend" / "artifacts" / "evaluation.json")
backtest = load_json(ROOT / "backend" / "artifacts" / "backtest.json")


def rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = rgb(color)
    if italic is not None:
        run.font.italic = italic


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color or COLORS["ink"])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        shade_cell(hdr[i], COLORS["light_blue"])
        set_cell_text(hdr[i], header, bold=True, color=COLORS["ink"], size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_kv_table(doc, rows, widths=(1.875, 4.625), fill=None):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, detail in rows:
        cells = table.add_row().cells
        if fill:
            shade_cell(cells[0], fill)
            shade_cell(cells[1], fill)
        set_cell_text(cells[0], label, bold=True, size=9.3)
        set_cell_text(cells[1], detail, size=9.3)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc, title, text, fill=COLORS["callout"], border_color=None):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(title + " ")
    set_run_font(r, size=10, bold=True, color=COLORS["dark_blue"])
    r = p.add_run(text)
    set_run_font(r, size=10, color=COLORS["ink"])
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    set_table_width(table, [6.5])
    doc.add_paragraph()
    return table


def add_p(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, size=11, bold=True, color=COLORS["ink"])
        run = p.add_run(text[len(bold_prefix):])
        set_run_font(run, size=11, color=COLORS["ink"])
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=COLORS["ink"])
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        run = p.add_run(item)
        set_run_font(run, size=10.7, color=COLORS["ink"])


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=10.7, color=COLORS["ink"])


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        set_run_font(run, size=16, bold=True, color=COLORS["blue"])
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    for run in p.runs:
        set_run_font(run, size=13, bold=True, color=COLORS["blue"])
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        set_run_font(run, size=12, bold=True, color=COLORS["dark_blue"])
    return p


def metric_value(metrics, market, key, default="n/d"):
    for item in metrics or []:
        if item.get("market") == market:
            value = item.get(key)
            if value is None:
                return default
            if isinstance(value, float):
                return f"{value:.4g}"
            return str(value)
    return default


def market_name(market):
    names = {
        "1X2": "1X2",
        "OVER_1_5_GOALS": "Over 1.5 gols",
        "OVER_2_5_GOALS": "Over 2.5 gols",
        "OVER_3_5_GOALS": "Over 3.5 gols",
        "UNDER_2_5_GOALS": "Under 2.5 gols",
        "UNDER_3_5_GOALS": "Under 3.5 gols",
        "BOTH_TEAMS_SCORE": "Ambas marcam",
        "DOUBLE_CHANCE": "Dupla chance",
        "CARDS": "Cartoes",
        "CORNERS": "Escanteios",
    }
    return names.get(market, market)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "BetIntel AI - guia de estudo para apresentacao"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=COLORS["muted"])

    footer = section.footer.paragraphs[0]
    footer.text = "Uso academico. Analise baseada em dados historicos. Nao garante resultado."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=8.5, color=COLORS["muted"])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("BetIntel AI")
    set_run_font(r, size=26, bold=True, color=COLORS["dark_blue"])
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Guia detalhado de estudo para apresentar o trabalho de Inteligencia Artificial")
    set_run_font(r, size=14, color=COLORS["muted"])

    add_kv_table(
        doc,
        [
            ("Projeto", "BetIntel AI - plataforma academica de analise probabilistica de futebol."),
            ("Data do guia", "25/06/2026, preparado para a apresentacao de 26/06/2026."),
            ("Entregavel do professor", "Solucao aplicada de IA + harness de desenvolvimento com specs, contexto, agents/skills, testes, evals e documentacao."),
            ("Mensagem central", "O projeto nao promete resultado: calcula estimativas educacionais com dados historicos e mostra quando faltam dados."),
        ],
        fill=COLORS["light_gray"],
    )

    add_callout(
        doc,
        "Frase segura para abrir a apresentacao:",
        "O BetIntel AI e uma solucao academica que usa machine learning probabilistico para estimar cenarios de partidas de futebol. Ele nao usa previsoes prontas da API, nao recomenda aposta financeira e sempre informa quando um mercado tem dados insuficientes.",
    )

    h1(doc, "1. Resumo em 60 segundos")
    add_p(
        doc,
        "O BetIntel AI resolve o problema de organizar dados historicos e atuais de futebol para gerar uma analise probabilistica educacional por partida. A solucao tem frontend em React, backend Node/TypeScript, ingestao de dados, cache local, treinamento, avaliacao, backtesting e um harness documentado para orientar o uso de IA no desenvolvimento.",
    )
    add_bullets(
        doc,
        [
            "A IA do produto nao e uma LLM conversando com o usuario: e um modelo supervisionado/probabilistico treinado com jogos historicos.",
            "A IA generativa foi usada como assistente de desenvolvimento no estilo Codex/Claude Code, seguindo arquivos de contexto, prompts, agentes, skills, specs e criterios de aceite.",
            "O modelo aprende frequencias historicas por competicao/temporada e perfis dos times, depois calcula probabilidades para mercados como 1X2, gols, ambas marcam, dupla chance, cartoes e escanteios.",
            "Quando faltam dados de cartoes ou escanteios, o backend retorna dados_insuficientes em vez de inventar uma probabilidade.",
            "A validacao inclui testes unitarios, build TypeScript, avaliacao com accuracy/Brier/cobertura e backtesting temporal.",
        ]
    )

    h1(doc, "2. O que o professor pediu no PDF")
    add_p(
        doc,
        "O enunciado exige dois elementos obrigatorios: uma solucao aplicada de IA e um harness de desenvolvimento. O professor tambem pede que o aluno demonstre dominio critico, explique as decisoes tecnicas, valide o funcionamento e documente como outra pessoa pode reproduzir ou continuar o projeto.",
    )
    add_table(
        doc,
        ["Pedido do PDF", "Como o BetIntel AI atende"],
        [
            ["Solucao aplicada de IA", "Sistema de analise probabilistica de futebol com modelo treinado, backend, frontend e pipeline de dados."],
            ["Specification-driven development", "Pasta specs/ com projeto, requisitos, criterios de aceite e mercados backend."],
            ["Harness/context engineering", "Arquivos .claude/context.md, prompts.md, agents e skills orientam a IA de desenvolvimento."],
            ["Agents e skills", "Agente analista-esportivo e skills de data-ingestion/model-evaluation documentadas."],
            ["Testes/evals", "backend/src/markets.test.ts e evals/backend-evals.md; comandos backend:test, evaluate e backtest."],
            ["Documentacao", "README.md e docs/ com arquitetura, fontes de dados, modelo ML e validacao."],
            ["Validacao tecnica", "Build, testes unitarios, avaliacao de metricas, backtesting temporal e retorno dados_insuficientes."],
            ["Dominio critico de IA", "O projeto declara limites, nao promete acuracia irreal, nao usa odds e diferencia estimativa educacional de recomendacao financeira."],
        ],
        [2.25, 4.25],
    )

    add_callout(
        doc,
        "Resposta direta:",
        "Sim, o projeto esta alinhado ao pedido do professor. A ressalva e que a apresentacao deve enfatizar os limites: o modelo atual e um baseline probabilistico auditavel, nao uma rede neural profunda nem uma garantia de acerto.",
        fill=COLORS["green_fill"],
    )

    h1(doc, "3. Qual IA estou usando?")
    h2(doc, "3.1 IA usada no desenvolvimento")
    add_p(
        doc,
        "No desenvolvimento, foi usada IA generativa como assistente de engenharia. Neste ambiente, o assistente e o Codex baseado em GPT-5; o repositorio tambem possui harness em .claude/ e CLAUDE.md porque a estrutura foi organizada no padrao Claude Code/agentic coding. Para o professor, a forma correta de explicar e: usei IA generativa para acelerar desenvolvimento, mas com especificacoes, contexto, testes e revisao critica.",
    )
    add_bullets(
        doc,
        [
            "A IA generativa ajudou a criar/ajustar codigo, documentacao, testes, providers, endpoints e material de estudo.",
            "O harness limita o comportamento da IA: nao prometer lucro, nao usar odds, nao inventar dados e retornar dados_insuficientes quando faltar informacao.",
            "O aluno continua responsavel por entender o codigo, rodar validacoes e explicar as decisoes.",
        ]
    )
    h2(doc, "3.2 IA usada dentro do produto")
    add_p(
        doc,
        "Dentro do BetIntel AI, a IA principal e um modelo de machine learning supervisionado/probabilistico. Ele aprende a partir de jogos historicos com placar, times, competicao, temporada, cartoes e escanteios quando disponiveis. O modelo nao e uma previsao pronta da API-Football, nao consulta Google para decidir o resultado e nao usa odds reais como recomendacao.",
    )
    add_kv_table(
        doc,
        [
            ("Tipo", "Baseline supervisionado de frequencias historicas segmentadas com ajuste por perfil de time."),
            ("Por que esse tipo?", "E auditavel, facil de explicar e adequado para demonstrar feature engineering, labels, avaliacao e backtesting."),
            ("Entrada", "Partida com mandante, visitante, liga/competicao, temporada e dados historicos do cache."),
            ("Saida", "Mercados disponiveis, mercados ignorados, probabilidades por selecao, sampleSize, confidence, fonte e updatedAt."),
            ("Limite", "Nao garante resultado e nao deve ser tratado como recomendacao financeira."),
        ],
    )

    h1(doc, "4. Arquitetura do sistema")
    add_p(
        doc,
        "A arquitetura foi separada para demonstrar engenharia de software e ML de ponta a ponta. O frontend mostra jogos e analises; o backend concentra dados, treino, avaliacao, backtest e predicao; o cache evita depender da API externa a cada render.",
    )
    add_table(
        doc,
        ["Camada", "Responsabilidade", "Arquivos principais"],
        [
            ["Frontend", "Exibe filtros, lista de partidas, painel de analise, mercados disponiveis/ignorados e aviso etico.", "src/App.tsx, src/components/*, src/lib/api.ts"],
            ["Backend HTTP", "Fornece endpoints /health, /markets, /competitions, /fixtures, /predict, /train, /evaluation e /backtest.", "backend/src/server.ts"],
            ["Providers", "Buscam e normalizam dados da API-Football, Football-Data.co.uk e calendario oficial da Copa 2026.", "backend/src/providers/*"],
            ["Feature engineering", "Transforma linhas brutas em labels e registros padronizados.", "backend/src/featureEngineering.ts, markets.ts"],
            ["Treinamento", "Treina mercados por segmento e perfis de times.", "backend/src/training.ts"],
            ["Predicao", "Escolhe segmento, ajusta por perfil do time e retorna resposta auditavel.", "backend/src/prediction.ts"],
            ["Validacao", "Calcula metricas e backtesting temporal.", "backend/src/evaluation.ts, backtesting.ts"],
            ["Harness", "Guia o comportamento da IA de desenvolvimento e documenta requisitos.", "specs/, docs/, .claude/, CLAUDE.md"],
        ],
        [1.1, 3.05, 2.35],
        font_size=8.8,
    )
    add_callout(
        doc,
        "Fluxo resumido:",
        "sync-data baixa/cacheia dados -> featureEngineering normaliza -> training gera model.json -> evaluation/backtesting medem qualidade -> server expõe endpoints -> frontend mostra analise educacional.",
    )

    h1(doc, "5. Fontes de dados")
    add_p(
        doc,
        "O projeto usa fontes concretas e documentadas. A regra central e: fonte externa fornece dados, nao a decisao da IA. A decisao do modelo vem do treino local.",
    )
    add_table(
        doc,
        ["Fonte", "Uso", "Observacao critica"],
        [
            ["API-Football / API-Sports", "Fixtures atuais e resultados historicos quando API_FOOTBALL_KEY permite acesso.", "Copa 2026 usa league=1 e season=2026; plano gratuito pode bloquear temporadas 2025/2026."],
            ["Football-Data.co.uk", "Historico CSV para ligas europeias, com placar e estatisticas quando disponiveis.", "Odds sao ignoradas. Cartoes/escanteios dependem das colunas disponiveis."],
            ["Calendario oficial Copa 2026", "Fallback real estatico para datas e confrontos definidos da Copa quando API nao libera 2026.", "Nao e placar ao vivo; serve para fixtures futuras."],
            ["Planilha das Copas 2006-2022", "Dataset de apoio gerado para demonstrar partidas, gols, cartoes e escanteios das ultimas cinco Copas concluidas.", "Pode ser importada como extensao do treino; nao deve ser confundida com previsao pronta."],
            ["Google", "Nao usado como dataset de treino.", "Google agrega dados de terceiros; raspagem automatizada nao e fonte reproduzivel/licenciada para treinar modelo."],
            ["Opta / Stats Perform", "Possivel melhoria futura com contrato/licenca.", "Sem contrato/API key/documentacao, nao deve ser copiado nem raspado."],
        ],
        [1.35, 2.45, 2.7],
        font_size=8.7,
    )
    add_p(
        doc,
        f"Planilha de apoio criada: {EXCEL}",
        bold_prefix="Planilha de apoio criada:",
    )

    h1(doc, "6. Como o modelo e treinado")
    add_numbers(
        doc,
        [
            "Sincronizacao: npm run backend:sync baixa dados atuais e historicos permitidos, depois grava backend/data/combined-results.csv, fixtures.json e sync-metadata.json.",
            "Normalizacao: providers convertem cada fonte para um schema interno comum, com campos como mandante, visitante, placar, liga, temporada, cartoes e escanteios.",
            "Feature engineering: totalGoals, outcome, totalCards e totalCorners sao derivados quando os campos existem.",
            "Labels: markets.ts cria labels de 1X2, over/under, ambas marcam, dupla chance, cartoes e escanteios.",
            "Treino por mercado: training.ts treina cada mercado separadamente e ignora mercado sem linhas validas suficientes.",
            "Segmentacao: o modelo calcula frequencias globais, por liga, por temporada, liga+temporada, competicao e competicao+temporada.",
            "Perfis dos times: o treino cria teamProfiles com vitorias, empates, derrotas, gols pro/contra, over/under, BTTS, cartoes e escanteios quando existem dados.",
            "Predicao: prediction.ts escolhe o melhor segmento disponivel e ajusta probabilidades usando o perfil do mandante e visitante.",
            "Saida auditavel: /predict retorna availableMarkets, ignoredMarkets, reason, sourceProvider, updatedAt, sampleSize e confidence.",
        ]
    )
    add_callout(
        doc,
        "Ponto para falar se perguntarem por que nao aparece tudo:",
        "Cartoes e escanteios nao existem em todos os datasets. O sistema foi projetado para nao quebrar e nao inventar dado: mercado sem amostra suficiente vira dados_insuficientes.",
        fill=COLORS["warn"],
    )

    h1(doc, "7. Features, labels e mercados")
    add_table(
        doc,
        ["Mercado", "Label usado no treino", "Campos de origem"],
        [
            ["1X2", "H se gols casa > gols fora; D se empate; A se fora > casa.", "FTHG, FTAG"],
            ["Over 1.5 / 2.5 / 3.5", "totalGoals > linha.", "FTHG + FTAG"],
            ["Under 2.5 / 3.5", "totalGoals < linha.", "FTHG + FTAG"],
            ["Ambas marcam", "FTHG > 0 e FTAG > 0.", "FTHG, FTAG"],
            ["Dupla chance", "1X = H ou D; 12 = H ou A; X2 = D ou A.", "Resultado 1X2"],
            ["Escanteios", "Over 8.5 e Over 9.5 escanteios.", "HC + AC"],
            ["Cartoes", "Over 3.5, 4.5 e 5.5 cartoes.", "HY + AY + HR + AR"],
        ],
        [1.55, 3.0, 1.95],
        font_size=8.8,
    )
    add_p(
        doc,
        "Essas labels sao importantes porque mostram que o modelo aprende de resultados observados. Exemplo: totalGoals e calculado; nao vem de uma previsao pronta da API.",
    )

    h1(doc, "8. Estado atual do modelo")
    providers = ", ".join(model.get("sourceProviders", [])) or "n/d"
    competitions = ", ".join(model.get("competitions", [])) or "n/d"
    team_profiles = len(model.get("teamProfiles", {}) or {})
    add_kv_table(
        doc,
        [
            ("Linhas no modelo atual", str(model.get("trainingRows", "n/d"))),
            ("Providers do modelo", providers),
            ("Competicoes no artefato", competitions),
            ("Perfis de time", str(team_profiles)),
            ("Minimo por mercado/segmento", str(model.get("minRows", "n/d"))),
            ("Criado em", str(model.get("createdAt", "n/d"))),
            ("Dados atualizados ate", str(model.get("updatedAt", "n/d"))),
        ],
    )
    market_rows = []
    for market, data in (model.get("markets") or {}).items():
        market_rows.append([
            market_name(market),
            data.get("status", "n/d"),
            str(data.get("usableRows", "n/d")),
            data.get("reason") or "ok",
        ])
    add_table(
        doc,
        ["Mercado", "Status", "Linhas validas", "Motivo/observacao"],
        market_rows,
        [1.45, 1.0, 1.1, 2.95],
        font_size=8.3,
    )

    h1(doc, "9. Avaliacao e backtesting")
    add_p(
        doc,
        "A avaliacao nao serve para vender certeza; serve para medir o comportamento do modelo. Accuracy mede a proporcao de acertos da selecao escolhida. Brier score mede qualidade de probabilidade: quanto menor, melhor. Cobertura mostra em que percentual de casos o mercado foi avaliavel.",
    )
    eval_metrics = evaluation.get("metrics", [])
    back_metrics = backtest.get("metrics", [])
    metric_rows = []
    for market in [
        "1X2",
        "OVER_1_5_GOALS",
        "OVER_2_5_GOALS",
        "OVER_3_5_GOALS",
        "BOTH_TEAMS_SCORE",
        "DOUBLE_CHANCE",
        "CARDS",
        "CORNERS",
    ]:
        metric_rows.append(
            [
                market_name(market),
                metric_value(eval_metrics, market, "selectionAccuracy"),
                metric_value(eval_metrics, market, "brierScore"),
                metric_value(eval_metrics, market, "coverage"),
                metric_value(back_metrics, market, "selectionAccuracy"),
                metric_value(back_metrics, market, "coverage"),
            ]
        )
    add_table(
        doc,
        ["Mercado", "Eval acc. %", "Eval Brier", "Eval cob. %", "Backtest acc. %", "Backtest cob. %"],
        metric_rows,
        [1.45, 0.9, 0.85, 0.85, 1.1, 1.35],
        font_size=8.2,
    )
    add_callout(
        doc,
        "Como interpretar:",
        "O objetivo academico nao e prometer 90% de acerto. Futebol e ruidoso, e mercados como 1X2, ambas marcam, cartoes e escanteios sao naturalmente incertos. Um resultado honesto com Brier/cobertura e melhor do que uma promessa sem validacao.",
    )

    h1(doc, "10. Frontend e experiencia do usuario")
    add_p(
        doc,
        "O frontend demonstra a solucao aplicada: carrega jogos futuros, mostra filtros, lista confrontos e abre o painel de analise. A Copa do Mundo 2026 aparece como competicao, e os jogos saem da lista quando chega o horario de inicio.",
    )
    add_bullets(
        doc,
        [
            "O App consulta o backend a cada 30 segundos.",
            "Tambem ha poda local a cada 5 segundos para remover jogos cujo horario ja passou.",
            "Se o backend estiver indisponivel, a interface mostra fallback mockado claramente marcado.",
            "O painel mostra competicao, data/hora, fonte, updatedAt, mercados disponiveis, mercados ignorados e aviso etico.",
            "Nao ha casas de aposta, logos nem odds reais.",
        ]
    )

    h1(doc, "11. Endpoints e comandos para demonstrar")
    add_table(
        doc,
        ["Endpoint", "Funcao"],
        [
            ["GET /health", "Verifica se o backend esta vivo."],
            ["GET /markets", "Lista mercados obrigatorios."],
            ["GET /competitions", "Lista competicoes cacheadas."],
            ["POST /sync-data", "Sincroniza dados externos/cache."],
            ["POST /train", "Treina e salva model.json."],
            ["GET /evaluation", "Retorna metricas de avaliacao."],
            ["GET /backtest", "Retorna backtesting temporal."],
            ["POST /predict", "Retorna probabilidades, mercados disponiveis/ignorados e aviso etico."],
            ["GET /fixtures?competition=&from=&to=", "Retorna jogos futuros filtrados."],
        ],
        [2.1, 4.4],
        font_size=8.7,
    )
    add_table(
        doc,
        ["Comando", "Quando usar na apresentacao"],
        [
            ["npm install", "Instalar dependencias."],
            ["npm run build", "Provar que frontend/backend compilam."],
            ["npm run backend:test", "Rodar testes unitarios."],
            ["npm run backend:sync", "Atualizar cache de dados."],
            ["npm run backend:train", "Treinar modelo."],
            ["npm run backend:evaluate", "Gerar avaliacao holdout."],
            ["npm run backend:backtest", "Gerar backtesting temporal."],
            ["npm run backend:serve", "Subir backend local."],
            ["npm run dev", "Subir frontend Vite."],
        ],
        [2.2, 4.3],
        font_size=8.7,
    )

    h1(doc, "12. Harness e SDD")
    add_p(
        doc,
        "O harness e a parte do trabalho que mostra como a IA foi orientada. Ele evita prompts soltos e cria um sistema de regras, contexto e validacao. Isso e exatamente o que o PDF pede quando fala em harness/context engineering.",
    )
    add_table(
        doc,
        ["Arquivo/pasta", "Papel no harness"],
        [
            ["specs/projeto.md", "Objetivo, problema, escopo e visao do projeto."],
            ["specs/requisitos.md", "Requisitos funcionais, nao funcionais e mercados obrigatorios."],
            ["specs/criterios-aceite.md", "Criterios que definem se a entrega esta aceitavel."],
            ["docs/arquitetura.md", "Arquitetura tecnica e fluxo de dados."],
            ["docs/fontes-dados.md", "Fontes concretas, limitacoes e regra de nao usar odds."],
            ["docs/modelo-ml.md", "Abordagem de ML, labels, segmentacao e metricas."],
            ["docs/validacao.md", "Como testar, avaliar e fazer backtest."],
            [".claude/context.md", "Contexto persistente para a IA de desenvolvimento."],
            [".claude/prompts.md", "Registro dos prompts principais."],
            [".claude/agents/...", "Agente analista-esportivo para orientar tarefas."],
            [".claude/skills/...", "Skills de ingestao e avaliacao."],
            ["CLAUDE.md / AGENTS.md", "Mapa rapido de regras para qualquer agente que alterar o projeto."],
            ["evals/backend-evals.md", "Registro de avaliacoes e validacao."],
        ],
        [2.3, 4.2],
        font_size=8.5,
    )

    h1(doc, "13. O que dizer se perguntarem sobre Google, Opta e 90%")
    add_kv_table(
        doc,
        [
            ("Google", "Nao e adequado raspar Google para treino. Ele mostra dados agregados de terceiros, mas nao e uma fonte licenciada e reproduzivel para dataset academico robusto."),
            ("Opta", "Seria uma fonte excelente para xG, eventos, cartoes, escanteios e dados granulares, mas exige contrato/licenca/API key. O projeto documenta como integrar no futuro sem scraping."),
            ("90% de acerto", "Nao e uma meta honesta para mercados reais de futebol. Um modelo pode atingir 90% em labels triviais/desbalanceadas, mas isso nao significa utilidade. O correto e apresentar accuracy, Brier, cobertura e limites."),
            ("Mais dados", "Mais dados ajudam, principalmente se tiverem qualidade e cobertura de cartoes/escanteios, mas precisam ser avaliados temporalmente para evitar overfitting."),
        ],
        widths=(1.4, 5.1),
        fill=COLORS["warn"],
    )

    h1(doc, "14. Perguntas provaveis do professor e respostas")
    qa_rows = [
        ["Qual e a IA do projeto?", "Um modelo probabilistico supervisionado, treinado com historico de partidas e ajustado por contexto/time. A IA generativa foi usada como assistente de desenvolvimento com harness."],
        ["Voce usou previsao pronta da API?", "Nao. A API fornece dados. A decisao probabilistica e calculada pelo modelo local treinado no backend."],
        ["Como o modelo aprende?", "Ele deriva labels dos resultados reais e calcula frequencias por mercado, segmento e perfil de time. Depois valida por holdout e backtest temporal."],
        ["Por que cartoes/escanteios aparecem como dados_insuficientes?", "Porque nem todo CSV/API tem essas colunas. O sistema prefere ignorar o mercado a inventar dado."],
        ["O sistema garante resultado?", "Nao. Ele mostra estimativas educacionais com aviso visivel e nao recomenda aposta financeira."],
        ["Como sei que nao e so mock?", "O backend possui providers reais, cache local, CLI de sync, endpoints, artifacts de treino/avaliacao e testes. Quando a API nao libera 2026, o fallback fica marcado."],
        ["Esta de acordo com o PDF?", "Sim: tem solucao aplicada de IA, SDD, harness, agentes/skills, testes/evals, docs, validacao e reproducao por README."],
        ["O que voce melhoraria?", "Integrar Opta/Sportmonks/Sportradar com licenca, criar features de xG, Elo, forma recente, descanso, lesoes e calibracao probabilistica."],
    ]
    add_table(doc, ["Pergunta", "Resposta curta"], qa_rows, [2.05, 4.45], font_size=8.5)

    h1(doc, "15. Roteiro de apresentacao")
    add_numbers(
        doc,
        [
            "Comece pelo problema: analises de futebol costumam misturar palpite, odds e promessa. O BetIntel organiza dados e retorna estimativas educacionais auditaveis.",
            "Mostre a arquitetura: frontend, backend, providers, cache, treino, avaliacao e predicao.",
            "Explique a IA: modelo probabilistico supervisionado com features/labels, nao previsao pronta.",
            "Abra a planilha/artefatos: mostre dados historicos, labels e fontes.",
            "Mostre o frontend: Copa 2026 no filtro, partidas futuras e painel de analise.",
            "Mostre mercados disponiveis e ignoredMarkets: esse e um ponto forte de engenharia, porque mostra dominio sobre dados faltantes.",
            "Mostre validacao: testes, build, evaluation e backtest.",
            "Feche com etica e limites: sem promessa de lucro, sem odds, sem Google scraping e sem 90% artificial.",
        ]
    )

    h1(doc, "16. Checklist final antes de entregar")
    add_bullets(
        doc,
        [
            "Repositorio no GitHub e link postado no Google Classroom ate 26/06/2026 ao meio-dia.",
            "README.md explica como rodar, configurar API_FOOTBALL_KEY, sincronizar, treinar, avaliar e fazer backtest.",
            "npm run build passa.",
            "npm run backend:test passa.",
            "Frontend carrega e mostra Copa do Mundo 2026.",
            "Backend /health responde.",
            "Painel mostra aviso: Analise baseada em dados historicos. Nao garante resultado.",
            "Nenhuma tela promete lucro, green garantido ou recomendacao financeira.",
            "Voce consegue explicar a diferenca entre IA generativa usada no desenvolvimento e modelo ML usado no produto.",
        ]
    )

    h1(doc, "17. Limites honestos e proximas melhorias")
    add_p(
        doc,
        "O projeto esta correto para o escopo academico, mas ainda e um baseline. Isso e uma escolha defensavel porque o trabalho valoriza especificacao, harness, validacao e dominio critico, nao apenas um modelo complexo sem explicacao.",
    )
    add_bullets(
        doc,
        [
            "A cobertura de cartoes e escanteios ainda depende da disponibilidade das fontes.",
            "A API-Football pode bloquear temporadas no plano gratuito; o fallback documentado evita quebrar o sistema.",
            "O modelo atual nao usa redes neurais profundas, xG, lesoes, escalações ou dados evento-a-evento.",
            "A planilha das Copas 2006-2022 deve ser apresentada como dataset historico de apoio/expansao, a menos que seja importada explicitamente no pipeline de treino.",
            "Melhorias futuras: Elo por time, forma recente, calibracao Platt/isotonic, validação por competição, provider licenciado Opta/Sportradar/Sportmonks e monitoramento de drift.",
        ]
    )

    doc.save(DOCX)
    return DOCX


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    print(build_document())
